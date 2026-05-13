"""Apply surgical edits to qavante_fase1_v2.6.4.docx:
  1) Header version: v2.6.3 -> v2.6.4
  2) Email notify@qavante.cl -> notify@qavante.com

Works at the run level so formatting is preserved. Cross-run matches use
full-paragraph rewrite as fallback (only when needed, only for paragraphs
where the substring spans runs).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from docx import Document

PATH = "qavante_fase1_v2.6.4.docx"

REPLACEMENTS = [
    ("Documento Maestro v2.6.3 para Claude Code", "Documento Maestro v2.6.4 para Claude Code"),
    ("notify@qavante.cl", "notify@qavante.com"),
]


def replace_in_paragraph(p, old, new):
    """Replace old with new inside paragraph p, preserving run formatting when possible."""
    if old not in p.text:
        return False
    # First try: substring lives entirely inside one run
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Fallback: collapse runs into the first one (loses inline formatting on this paragraph only)
    full = p.text.replace(old, new)
    if not p.runs:
        p.text = full
        return True
    p.runs[0].text = full
    for r in p.runs[1:]:
        r.text = ""
    return True


doc = Document(PATH)
changes = {old: 0 for old, _ in REPLACEMENTS}

def walk(paragraphs):
    for p in paragraphs:
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                changes[old] += 1

walk(doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            walk(cell.paragraphs)

doc.save(PATH)

print("Cambios aplicados:")
for old, count in changes.items():
    print(f"  {count}x  {old!r}")
