"""Read-only inventory of qavante.cl / qavante.com / NIC Chile mentions in the master doc."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from docx import Document

PATH = "qavante_fase1_v2.6.4.docx"
NEEDLES = [
    "qavante.cl", "qavante.com", "qavante.",
    "NIC Chile", "nic.cl",
    "https://", "http://",
    "@qavante", "@tooxs",
    "/login", "/aceptar-invitacion",
    "dominio",
]

doc = Document(PATH)

print(f"Archivo: {PATH}")
print(f"Paragraphs total: {len(doc.paragraphs)}")
print(f"Tables total: {len(doc.tables)}")
print()

hits = []

for i, p in enumerate(doc.paragraphs):
    txt = p.text
    if any(n.lower() in txt.lower() for n in NEEDLES):
        style = p.style.name if p.style is not None else "(no-style)"
        hits.append(("P", i, style, txt))

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for pi, p in enumerate(cell.paragraphs):
                txt = p.text
                if any(n.lower() in txt.lower() for n in NEEDLES):
                    style = p.style.name if p.style is not None else "(no-style)"
                    hits.append(("T", f"t{ti}r{ri}c{ci}p{pi}", style, txt))

print(f"Ocurrencias encontradas: {len(hits)}")
print("=" * 80)
for kind, loc, style, txt in hits:
    print(f"[{kind} {loc}] style={style}")
    print(f"  > {txt}")
    print()

# Version sniff
print("=" * 80)
print("PRIMER PARRAFO (header probable):")
for i in range(min(5, len(doc.paragraphs))):
    print(f"  [{i}] {doc.paragraphs[i].text!r}")
